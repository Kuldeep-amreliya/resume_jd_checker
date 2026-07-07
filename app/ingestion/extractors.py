"""
Format-specific text extraction.

Each function takes raw file bytes and returns plain text plus any warnings.
None of these raise for "this looks empty" — they report it via warnings/
low_text_confidence and let the router decide what to do. They DO raise
IngestionError for genuinely corrupt/unreadable files.
"""
from app.ingestion.ingestion_schemas import IngestionError


def extract_pdf_text(file_bytes: bytes, low_text_threshold: int) -> tuple[str, bool, list[str]]:
    """
    Extract text from a PDF using PyMuPDF (fitz).

    Returns (text, low_text_confidence, warnings).
    low_text_confidence is True when the PDF likely has no real text layer
    (e.g. it's a scanned image) — total extracted chars fall below threshold
    despite the PDF having pages, which usually means OCR would be needed.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise IngestionError(
            "PDF support is not installed on the server.",
            detail="Install with: pip install pymupdf",
        ) from e

    warnings: list[str] = []

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        raise IngestionError("Could not open PDF — the file may be corrupt.", detail=str(e)) from e

    if doc.page_count == 0:
        doc.close()
        raise IngestionError("PDF has no pages.")

    pages_text = []
    for page in doc:
        pages_text.append(page.get_text("text"))
    doc.close()

    text = "\n".join(pages_text).strip()

    low_text_confidence = len(text) < low_text_threshold
    if low_text_confidence:
        warnings.append(
            "Very little text could be extracted from this PDF. "
            "It may be a scanned/image-based document that requires OCR."
        )

    return text, low_text_confidence, warnings


def extract_docx_text(file_bytes: bytes) -> tuple[str, list[str]]:
    """
    Extract text from a .docx using python-docx.

    Reads both paragraphs and table cells — skills/education are sometimes
    laid out in tables rather than prose, and we don't want to silently drop them.
    """
    import io

    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise IngestionError(
            "DOCX support is not installed on the server.",
            detail="Install with: pip install python-docx",
        ) from e

    warnings: list[str] = []

    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise IngestionError("Could not open DOCX — the file may be corrupt or not a valid .docx.", detail=str(e)) from e

    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    text = "\n".join(parts).strip()

    if not text:
        warnings.append("No readable text found in this DOCX file.")

    return text, warnings


def extract_txt_text(file_bytes: bytes) -> tuple[str, list[str]]:
    """
    Decode a .txt file. Tries UTF-8 first, falls back to chardet-detected
    encoding, then finally latin-1 (which never fails to decode, as a last resort).
    """
    warnings: list[str] = []

    try:
        return file_bytes.decode("utf-8").strip(), warnings
    except UnicodeDecodeError:
        pass

    try:
        import chardet

        detected = chardet.detect(file_bytes)
        encoding = detected.get("encoding")
        confidence = detected.get("confidence", 0)
        if encoding:
            warnings.append(
                f"File was not valid UTF-8; decoded as {encoding} (confidence {confidence:.0%})."
            )
            return file_bytes.decode(encoding, errors="replace").strip(), warnings
    except ImportError:
        pass

    warnings.append("Could not reliably detect file encoding; some characters may be incorrect.")
    return file_bytes.decode("latin-1", errors="replace").strip(), warnings